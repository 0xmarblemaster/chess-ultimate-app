"""
Extract Service

Service for extracting content from various document types.
"""

import os
import json
import re
from typing import List, Dict, Any, Tuple, Optional
import docx
import fitz  # PyMuPDF
from PIL import Image
import io
import logging
import traceback

# Import configuration
from backend import config

logger = logging.getLogger(__name__)

class ExtractService:
    """
    Service for extracting content from document files.
    
    This service handles extraction of text and images from:
    - DOCX files
    - PDF files
    
    It processes the documents to identify:
    - Lessons
    - Tasks
    - Diagrams/Chess positions
    - Explanations
    """
    
    def __init__(self, output_images_dir: Optional[str] = None):
        """
        Initialize the Extract service.
        
        Args:
            output_images_dir: Directory to store extracted images
        """
        self.logger = logger
        self.output_images_dir = output_images_dir or config.OUTPUT_IMAGE_DIR
        os.makedirs(self.output_images_dir, exist_ok=True)
        
        # Patterns for image extraction
        self.image_number_pattern = re.compile(r"(?:diagram|img)(?:_doc|_pdf|_page\d+)?(?:_img)?(?:_)?(\d+)(?:\.|_)?", re.IGNORECASE)
        self.any_number_pattern = re.compile(r"(\d+)", re.IGNORECASE)

    def sanitize_filename(self, name: str) -> str:
        """
        Sanitizes a string to be a valid filename.
        
        Args:
            name: The string to sanitize
            
        Returns:
            Sanitized filename
        """
        name = re.sub(r'[\/*?:"<>|]', "", name)  # Remove invalid characters
        name = name.replace(" ", "_") # Replace spaces with underscores
        return name[:200] # Limit length
        
    def extract_from_docx(self, file_path: str) -> Tuple[Dict[str, Any], List[Dict[str, Any]]]:
        """
        Extracts text and images from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Tuple of (structured data, extracted elements)
        """
        try:
            doc = docx.Document(file_path)
            self.logger.info(f"Successfully opened DOCX file: {file_path}")
            self.logger.info(f"Document has {len(doc.paragraphs)} paragraphs")
        except Exception as e:
            self.logger.error(f"Error opening DOCX file: {e}")
            return {"book_title": "Error", "lessons": []}, []
            
        data = {"book_title": "Unknown Document", "lessons": []}
        current_lesson = None
        image_counter = 1
        extracted_elements = []
        task_count = 0
        
        # Track if we're in a task section to improve task identification
        in_task_section = False
        
        # Simplistic title extraction (first non-empty paragraph, or filename)
        for para in doc.paragraphs:
            if para.text.strip():
                data["book_title"] = para.text.strip()
                break
        if data["book_title"] == "Unknown Document":
            data["book_title"] = os.path.splitext(os.path.basename(file_path))[0]

        # First pass - collect diagram references from text to help align image numbers
        diagram_refs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Detect diagram references
                diagram_matches = re.findall(r"(диаграмма|diagram)\s*(\d+)", text, re.IGNORECASE)
                if diagram_matches:
                    for _, num in diagram_matches:
                        diagram_refs.append(int(num))

        if diagram_refs:
            self.logger.debug(f"Found diagram references: {diagram_refs}")
            
        # Main extraction loop
        for para in doc.paragraphs:
            text = para.text.strip()
            style = para.style.name.lower() if hasattr(para, 'style') and hasattr(para.style, 'name') else "normal"

            # Detect lesson titles (heuristic)
            if style.startswith("heading 1") or text.lower().startswith("урок") or text.lower().startswith("lesson"):
                if current_lesson:
                    extracted_elements.append(current_lesson)
                lesson_title_match = re.match(r"(урок|lesson)\s*(\d+)\s*[:.-]?(.*)", text, re.IGNORECASE)
                if lesson_title_match:
                    lesson_number = int(lesson_title_match.group(2))
                    lesson_title = lesson_title_match.group(3).strip() or f"Lesson {lesson_number}"
                else:
                    lesson_number = len(extracted_elements) + 1 # Fallback lesson number
                    lesson_title = text
                current_lesson = {"lesson_number": lesson_number, "title": lesson_title, "content": []}
                self.logger.debug(f"Detected lesson: {lesson_title} (#{lesson_number})")
            
            # Skip empty paragraphs
            if not text:
                continue
                
            # Add text content to the current lesson
            if current_lesson:
                # Detect if this is a task by looking for indicators
                content_item = {"type": "explanation", "text": text}
                
                # Check for diagram references in the text
                diagram_match = re.search(r"(диаграмма|diagram)\s*(\d+)", text, re.IGNORECASE)
                if diagram_match:
                    content_item["type"] = "task"
                    content_item["diagram_number_reference"] = int(diagram_match.group(2))
                    task_count += 1
                    in_task_section = True
                    self.logger.debug(f"Found task with diagram #{diagram_match.group(2)}")
                elif in_task_section and any(kw in text.lower() for kw in ["задание", "task", "exercise", "упражнение"]):
                    content_item["type"] = "task"
                    task_count += 1
                    self.logger.debug(f"Found task #{task_count} without explicit diagram reference")
                
                current_lesson["content"].append(content_item)
            else:
                # Content before any lesson detected - create a default lesson
                if not extracted_elements:
                    default_lesson_title = data["book_title"] + " - Introduction"
                    current_lesson = {"lesson_number": 0, "title": default_lesson_title, "content": []}
                    current_lesson["content"].append({"type": "explanation", "text": text})
                    self.logger.debug(f"Created default lesson: {default_lesson_title}")

        # Add the last lesson if it exists
        if current_lesson and current_lesson not in extracted_elements:
            extracted_elements.append(current_lesson)
            
        # Count images in the document
        img_count = 0
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                img_count += 1
                
        self.logger.debug(f"Document has {img_count} image relations")
        
        # Start image counter at 1 if no diagram references found
        next_image_number = 1
        img_idx = 0
        
        # Extract and process images
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_bytes = rel.target_part.blob
                try:
                    img = Image.open(io.BytesIO(image_bytes))
                    
                    # If we can align with known diagram numbers, use those
                    # Otherwise use sequential numbering
                    if img_idx < len(diagram_refs):
                        # Use diagrams refs for image numbers to align with tasks
                        image_number = diagram_refs[img_idx]
                    else:
                        # If we've gone through all known diagram refs, use sequential numbering 
                        # starting after the highest known diagram number
                        image_number = next_image_number
                        next_image_number += 1
                    
                    # Create filename with the aligned image number
                    img_filename = f"diagram_doc_img{image_number}.{img.format.lower() or 'png'}"
                    img_path = os.path.join(self.output_images_dir, img_filename)
                    img.save(img_path)
                    img_idx += 1
                    self.logger.debug(f"Saved image {img_idx}/{img_count}: {img_filename} with number {image_number}")
                    
                    # Try to find the corresponding task and attach the image to it
                    image_added = False
                    for lesson in extracted_elements:
                        for item in lesson.get("content", []):
                            if item.get("type") == "task" and item.get("diagram_number_reference") == image_number:
                                item["image"] = img_filename
                                image_added = True
                                break
                        if image_added:
                            break
                except Exception as e:
                    self.logger.error(f"Error processing image: {e}")
            
        # Populate the lessons list in the data structure
        data["lessons"] = extracted_elements

        return data, extracted_elements

    def extract_from_pdf(self, file_path: str) -> Tuple[Dict[str, Any], List[Dict[str, Any]]]:
        """
        Extracts text and images from a PDF file.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Tuple of (structured data, extracted elements)
        """
        try:
            doc = fitz.open(file_path)
            self.logger.info(f"Successfully opened PDF file: {file_path}")
            self.logger.info(f"Document has {len(doc)} pages")
        except Exception as e:
            self.logger.error(f"Error opening PDF file: {e}")
            return {"book_title": "Error", "lessons": []}, []
        
        data = {"book_title": "Unknown PDF Document", "lessons": []}
        current_lesson = None
        image_counter = 1  # Unique counter for PDF images
        task_count = 0
        
        # Track if we're in a task section to improve task identification
        in_task_section = False
        
        # Extract metadata for title if available
        if doc.metadata and doc.metadata.get("title"):
            data["book_title"] = doc.metadata["title"]
        else:
            data["book_title"] = os.path.splitext(os.path.basename(file_path))[0]
    
        # Consolidate extracted elements (lessons)
        extracted_lessons = []
    
        # First pass to collect diagram references
        diagram_refs = []
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text = page.get_text("text")
            diagram_matches = re.findall(r"(диаграмма|diagram)\s*(\d+)", text, re.IGNORECASE)
            for _, num in diagram_matches:
                diagram_refs.append(int(num))
            
        if diagram_refs:
            self.logger.debug(f"Found PDF diagram references: {diagram_refs}")
            
        # Main extraction loop
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text_blocks = page.get_text("dict")["blocks"]
            
            # Process text blocks
            for block in text_blocks:
                if block.get("type") == 0:  # Text block
                    for line in block.get("lines", []):
                        line_text = ""
                        for span in line.get("spans", []):
                            text = span.get("text", "").strip()
                            if text:
                                line_text += text + " "
                                
                        line_text = line_text.strip()
                        if not line_text:
                            continue
                            
                        # Check for lesson headers
                        lesson_match = re.match(r"(урок|lesson)\s*(\d+)\s*[:.-]?(.*)", line_text, re.IGNORECASE)
                        if lesson_match:
                            if current_lesson:
                                extracted_lessons.append(current_lesson)
                                
                            lesson_number = int(lesson_match.group(2))
                            lesson_title = lesson_match.group(3).strip() or f"Lesson {lesson_number}"
                            current_lesson = {"lesson_number": lesson_number, "title": lesson_title, "content": []}
                            self.logger.debug(f"Detected PDF lesson: {lesson_title} (#{lesson_number})")
                            continue
                            
                        # Check for headings that might be lesson titles
                        font_size = max([span.get("size", 0) for span in line.get("spans", [])])
                        is_bold = any([span.get("font", "").lower().find("bold") != -1 for span in line.get("spans", [])])
                        
                        if (font_size > 12 and is_bold) or any(heading in line_text.lower() for heading in ["chapter", "section", "part", "глава", "раздел", "часть"]):
                            if current_lesson:
                                extracted_lessons.append(current_lesson)
                            lesson_number = len(extracted_lessons) + 1
                            current_lesson = {"lesson_number": lesson_number, "title": line_text, "content": []}
                            self.logger.debug(f"Detected PDF heading as lesson: {line_text}")
                            continue
    
                        if line_text:
                            content_item = {"type": "explanation", "text": line_text}
                            diagram_match = re.search(r"(диаграмма|diagram)\s*(\d+)", line_text, re.IGNORECASE)
                            if diagram_match:
                                content_item["type"] = "task"
                                content_item["diagram_number_reference"] = int(diagram_match.group(2))
                                self.logger.debug(f"Found PDF task with diagram #{diagram_match.group(2)}")
                            
                            if current_lesson:
                                current_lesson["content"].append(content_item)
                            else: # Content before any lesson detected
                                if not extracted_lessons:
                                    default_lesson_title = data["book_title"] + " - Introduction"
                                    current_lesson = {"lesson_number": 0, "title": default_lesson_title, "content": []}
                                    extracted_lessons.append(current_lesson)
                                current_lesson["content"].append(content_item)
            
            # Extract images from the page
            img_list = page.get_images(full=True)
            
            if not img_list:
                continue
                
            self.logger.debug(f"Found {len(img_list)} images on page {page_num+1}")
                
            for img_idx, img_info in enumerate(img_list):
                try:
                    xref = img_info[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    img_ext = base_image["ext"]
                    
                    # Determine which image number to use
                    if img_idx < len(diagram_refs):
                        image_number = diagram_refs[img_idx]
                    else:
                        # Use sequential numbering if we run out of diagram refs
                        image_number = image_counter
                        image_counter += 1
                    
                    # Save the image
                    img_filename = f"diagram_pdf_page{page_num+1}_img{image_number}.{img_ext}"
                    img_path = os.path.join(self.output_images_dir, img_filename)
                    
                    with open(img_path, "wb") as img_file:
                        img_file.write(image_bytes)
                    
                    self.logger.debug(f"Saved PDF image from page {page_num+1}: {img_filename}")
                    
                    # Try to find corresponding task
                    image_added = False
                    for lesson in extracted_lessons:
                        for item in lesson.get("content", []):
                            if item.get("type") == "task" and item.get("diagram_number_reference") == image_number:
                                item["image"] = img_filename
                                image_added = True
                                break
                        if image_added:
                            break
                            
                except Exception as e:
                    self.logger.error(f"Error extracting image: {e}")
        
        # Add the last lesson if it exists
        if current_lesson and current_lesson not in extracted_lessons:
            extracted_lessons.append(current_lesson)
            
        # Populate the lessons list in the data structure
        data["lessons"] = extracted_lessons
    
        return data, extracted_lessons

    def extract_content(self, file_path: str, output_images_dir: Optional[str] = None) -> Dict[str, Any]:
        """
        Main function to extract content from a document file.
        Detects file type and calls appropriate extraction function.
        
        Args:
            file_path: Path to the document file (DOCX or PDF)
            output_images_dir: Path to directory for saving extracted images
        
        Returns:
            Dictionary with extracted content
        """
        if output_images_dir:
            self.output_images_dir = output_images_dir
            os.makedirs(self.output_images_dir, exist_ok=True)
        
        # Check if file exists
        if not os.path.isfile(file_path):
            abs_path = os.path.abspath(file_path)
            self.logger.error(f"File not found: {abs_path}")
            return None
        
        # If file path is relative to INPUT_DIR in config, make it absolute
        if not os.path.isabs(file_path):
            file_path = os.path.join(config.INPUT_DIR, file_path)
        
        # Get lowercase file extension
        _, file_ext = os.path.splitext(file_path)
        file_ext = file_ext.lower()
        
        try:
            if file_ext == '.docx':
                self.logger.info(f"Extracting content from DOCX file: {file_path}")
                content_data, _ = self.extract_from_docx(file_path)
                self.logger.info(f"DOCX extraction completed: {len(content_data.get('lessons', []))} lessons found")
                return content_data
            elif file_ext == '.pdf':
                self.logger.info(f"Extracting content from PDF file: {file_path}")
                content_data, _ = self.extract_from_pdf(file_path)
                self.logger.info(f"PDF extraction completed: {len(content_data.get('lessons', []))} lessons found")
                return content_data
            else:
                self.logger.error(f"Unsupported file format: {file_ext}. Please use DOCX or PDF files.")
                return None
        except Exception as e:
            self.logger.error(f"Error during extraction: {e}")
            self.logger.error(traceback.format_exc())
            return None
    
    def healthcheck(self) -> bool:
        """
        Perform a health check on the service.
        
        Returns:
            True if service is healthy
        """
        try:
            # Check if output directory exists and is writable
            test_file = os.path.join(self.output_images_dir, "healthcheck.txt")
            with open(test_file, "w") as f:
                f.write("healthcheck")
            os.remove(test_file)
            return True
        except Exception as e:
            self.logger.error(f"Health check failed: {e}")
            return False


if __name__ == "__main__":
    # Example usage
    logging.basicConfig(level=logging.INFO)
    service = ExtractService()
    if service.healthcheck():
        print("ExtractService is healthy!")
        
        # Example extraction
        test_file = os.getenv("TEST_DOCX_FILE", "example.docx")
        if os.path.exists(test_file):
            result = service.extract_content(test_file)
            if result:
                print(f"Successfully extracted content from {test_file}")
                print(f"Title: {result['book_title']}")
                print(f"Found {len(result['lessons'])} lessons")
                
                # Print first lesson content sample
                if result['lessons']:
                    first_lesson = result['lessons'][0]
                    print(f"\nSample - Lesson {first_lesson['lesson_number']}: {first_lesson['title']}")
                    print(f"Content items: {len(first_lesson['content'])}")
                    if first_lesson['content']:
                        print(f"First content item: {first_lesson['content'][0]['type']} - " 
                              f"{first_lesson['content'][0]['text'][:100]}...")
            else:
                print(f"Failed to extract content from {test_file}")
        else:
            print(f"Test file {test_file} not found. Set TEST_DOCX_FILE environment variable to test extraction.")
    else:
        print("ExtractService health check failed!") 