import os
import json
import re
from typing import List, Dict, Any, Tuple, Optional
try:
    import docx
except ImportError:
    try:
        # Alternative import if python-docx is installed differently
        from docx import Document
        class DocxWrapper:
            @staticmethod
            def Document(*args, **kwargs):
                return Document(*args, **kwargs)
        docx = DocxWrapper
    except ImportError:
        print("ERROR: Neither python-docx nor docx modules could be imported. Please install with 'pip install python-docx'")
import fitz  # PyMuPDF
from PIL import Image
import io
from . import config

# ADD: Import the enhanced diagram processor
try:
    from .enhanced_diagram_processor import apply_enhanced_diagram_associations
    ENHANCED_DIAGRAM_PROCESSOR_AVAILABLE = True
except ImportError:
    ENHANCED_DIAGRAM_PROCESSOR_AVAILABLE = False
    print("Enhanced diagram processor not available, using legacy association logic")

# Pattern for multi-number detection in a single paragraph
MULTI_NUMBER_PATTERN = re.compile(r'\b(\d{1,3})\b', re.IGNORECASE)
# Pattern for matching diagram tasks with numbers and underscores (single task format)
TASK_PATTERN = re.compile(r"^\s*(\d{1,3})\s*[_]+\s*$", re.IGNORECASE)
# Pattern for task numbers with underscores that might be separated
TASK_PATTERN_SPACED = re.compile(r"^\s*(\d{1,3})\s+[_]+\s*$", re.IGNORECASE)
# Pattern for matching dual task format: "25 _______  26 _______"
DUAL_TASK_PATTERN = re.compile(r"^\s*(\d{1,3})\s*[_]+\s+(\d{1,3})\s*[_]+\s*$", re.IGNORECASE)
# Pattern that matches just numbers that could be task numbers
NUMBER_PATTERN = re.compile(r"^\s*(\d{1,3})\s*$", re.IGNORECASE)
# Pattern for detecting diagram references
DIAGRAM_PATTERN = re.compile(r"(диаграмма|diagram)\s*(\d+)", re.IGNORECASE)
# Pattern for sections that contain tasks
TASK_SECTION_PATTERN = re.compile(r"домашнее\s+задание|задача|задачи|мат\s+в\s+\d+\s+ход", re.IGNORECASE)
# Pattern for extracting numbers from ANY image filename format
IMAGE_NUMBER_PATTERN = re.compile(r"(?:diagram|img)(?:_doc|_pdf|_page\d+)?(?:_img)?(?:_)?(\d+)", re.IGNORECASE)
# Secondary pattern for any number in filename as fallback
ANY_NUMBER_PATTERN = re.compile(r"(\d+)", re.IGNORECASE)

def sanitize_filename(name: str) -> str:
    """Sanitizes a string to be a valid filename."""
    name = re.sub(r'[\/*?:"<>|]', "", name)  # Remove invalid characters
    name = name.replace(" ", "_")  # Replace spaces with underscores
    return name[:200]  # Limit length

def find_task_by_number_across_all_lessons(lessons: List[Dict[str, Any]], task_number: int) -> Optional[Dict[str, Any]]:
    """
    Finds a task with the exact matching number across all lessons.
    
    Args:
        lessons: List of lesson data
        task_number: The task number to find
        
    Returns:
        The task item with matching number or None if not found
    """
    # First try: exact match by diagram_number_reference
    for lesson in lessons:
        for item in lesson.get("content", []):
            if (item.get("type") == "task" or item.get("type") == "general_task") and "image" not in item:
                if item.get("diagram_number_reference") == task_number:
                    return item
    
    # Second try: extract number from task text if it's a numbered task (e.g., "43 ___________")
    for lesson in lessons:
        for item in lesson.get("content", []):
            if (item.get("type") == "task" or item.get("type") == "general_task") and "image" not in item:
                text = item.get("text", "")
                task_match = TASK_PATTERN.match(text) or TASK_PATTERN_SPACED.match(text) or NUMBER_PATTERN.match(text)
                if task_match and int(task_match.group(1)) == task_number:
                    # Set the diagram number reference if not already set
                    if "diagram_number_reference" not in item:
                        item["diagram_number_reference"] = task_number
                    return item
    
    # No exact match found
    return None

def find_closest_task_by_number(lesson_data: Dict[str, Any], image_number: int) -> Optional[Dict[str, Any]]:
    """
    Finds the closest task to the given image number in a lesson.
    
    Args:
        lesson_data: The lesson data containing content items
        image_number: The number of the image to match with a task
        
    Returns:
        The closest task item or None if no suitable task found
    """
    closest_task = None
    min_distance = float('inf')
    
    # First pass: look for exact matching diagram number
    for item in lesson_data.get("content", []):
        if (item.get("type") == "task" or item.get("type") == "general_task") and "image" not in item:
            if item.get("diagram_number_reference") == image_number:
                return item  # Exact match
    
    # Second pass: try to extract number from task text and match
    for item in lesson_data.get("content", []):
        if (item.get("type") == "task" or item.get("type") == "general_task") and "image" not in item:
            text = item.get("text", "")
            task_match = TASK_PATTERN.match(text) or TASK_PATTERN_SPACED.match(text) or NUMBER_PATTERN.match(text)
            if task_match and int(task_match.group(1)) == image_number:
                # Set the diagram number reference if not already set
                if "diagram_number_reference" not in item:
                    item["diagram_number_reference"] = image_number
                return item
    
    # Third pass: find closest task by number 
    for item in lesson_data.get("content", []):
        if (item.get("type") == "task" or item.get("type") == "general_task") and "image" not in item:
            task_number = None
            
            # Try to get diagram number from reference
            if "diagram_number_reference" in item:
                task_number = item["diagram_number_reference"]
            # Try to extract from text if available
            elif "text" in item:
                text = item.get("text", "")
                task_match = TASK_PATTERN.match(text) or TASK_PATTERN_SPACED.match(text) or NUMBER_PATTERN.match(text)
                if task_match:
                    task_number = int(task_match.group(1))
                    
            if task_number is not None:
                distance = abs(task_number - image_number)
                if distance < min_distance:
                    min_distance = distance
                    closest_task = item
    
    # Only return if distance is reasonable (even more forgiving)
    if min_distance <= 30:
        # Set the diagram number reference if not already set
        if closest_task and "diagram_number_reference" not in closest_task:
            closest_task["diagram_number_reference"] = image_number
        return closest_task
    
    # Fourth pass: any task without an image in this lesson
    for item in lesson_data.get("content", []):
        if (item.get("type") == "task" or item.get("type") == "general_task") and "image" not in item:
            # Assign the diagram number to this task since we're matching it
            item["diagram_number_reference"] = image_number
            return item
                
    return None

def associate_image_with_task(lessons: List[Dict[str, Any]], img_filename: str, image_number: Optional[int] = None) -> bool:
    """
    Associates an image with the most appropriate task across all lessons.
    
    Args:
        lessons: List of lesson data
        img_filename: Filename of the image to associate
        image_number: Optional image number extracted from filename
        
    Returns:
        True if association was successful, False otherwise
    """
    if image_number is None:
        # First try: Use improved pattern to extract image number from filename
        match = IMAGE_NUMBER_PATTERN.search(img_filename)
        if match:
            try:
                image_number = int(match.group(1))
                print(f"DEBUG: Extracted image number {image_number} from filename {img_filename} using primary pattern")
            except (ValueError, IndexError):
                image_number = None
        
        # Second try: extract any number from filename
        if image_number is None:
            numbers = re.findall(ANY_NUMBER_PATTERN, img_filename)
            if numbers:
                try:
                    image_number = int(numbers[-1])  # Use last number in filename
                    print(f"DEBUG: Extracted image number {image_number} from filename {img_filename} using fallback pattern")
                except (ValueError, IndexError):
                    return False  # Invalid number format
            else:
                return False  # No number found in filename
    
    print(f"DEBUG: Trying to associate image {img_filename} with task number {image_number}")
    
    # Strategy 1: Find task by exact number match across ALL lessons
    matching_task = find_task_by_number_across_all_lessons(lessons, image_number)
    if matching_task:
        matching_task["image"] = img_filename
        print(f"DEBUG: Associated image {img_filename} with task by direct number match (diagram #{image_number}) across all lessons")
        return True
    
    # Strategy 2: Search from recent to earlier lessons for closest match
    for lesson in reversed(lessons):  # Start from most recent lesson
        closest_task = find_closest_task_by_number(lesson, image_number)
        if closest_task:
            closest_task["image"] = img_filename
            task_number = closest_task.get("diagram_number_reference")
            print(f"DEBUG: Associated image {img_filename} with task by closest number (diagram #{task_number}) in lesson {lesson.get('lesson_number')}")
            return True
    
    # Strategy 3: Create a new task if nothing found and number is valid
    if lessons and image_number > 0:
        # Find the most likely lesson to add the new task to
        # First try to find lesson 15 since that seems to be the one with proper diagram tasks
        target_lesson = None
        for lesson in lessons:
            if lesson.get("lesson_number") == 15:
                target_lesson = lesson
                break
                
        # If no lesson 15, use most recent lesson
        if not target_lesson and lessons:
            target_lesson = lessons[-1]
        
        if target_lesson:
            # Create a new task using the image number
            new_task = {
                "type": "general_task",
                "text": f"{image_number} _______________",
                "diagram_number_reference": image_number,
                "image": img_filename
            }
            
            target_lesson["content"].append(new_task)
            print(f"DEBUG: Created new task for image {img_filename} with diagram #{image_number} in lesson {target_lesson.get('lesson_number')}")
            return True
    
    print(f"DEBUG: Failed to associate image {img_filename} with any task")
    return False

def extract_text_and_images_from_docx(file_path: str, output_images_dir: str) -> Tuple[Dict[str, Any], List[Dict[str, Any]]]:
    """Extracts text and images from a DOCX file."""
    try:
        doc = docx.Document(file_path)
        print(f"Successfully opened DOCX file: {file_path}")
        print(f"Document has {len(doc.paragraphs)} paragraphs")
    except Exception as e:
        print(f"Error opening DOCX file: {e}")
        return {"book_title": "Error", "lessons": []}, []
        
    data = {"book_title": "Unknown Document", "lessons": []}
    current_lesson = None
    image_counter = 1
    extracted_elements = []
    task_count = 0
    
    # Track if we're in a task section to improve task identification
    in_task_section = False
    
    # Simplistic title extraction (first non-empty paragraph, or filename)
    for para in doc.paragraphs:
        if para.text.strip():
            data["book_title"] = para.text.strip()
            break
    if data["book_title"] == "Unknown Document":
        data["book_title"] = os.path.splitext(os.path.basename(file_path))[0]

    # Dump first 20 paragraphs to understand format better
    for i, para in enumerate(doc.paragraphs[:20]):
        text = para.text.strip()
        if text:
            print(f"DEBUG: Paragraph {i}: '{text}' (Style: {para.style.name})")

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
            
        style = para.style.name.lower()

        # Detect task sections
        if TASK_SECTION_PATTERN.search(text):
            in_task_section = True
            print(f"DEBUG: Entering task section at paragraph {i}: '{text}'")
        
        # Detect lesson titles (heuristic)
        if style.startswith("heading 1") or text.lower().startswith("урок") or text.lower().startswith("lesson"):
            in_task_section = False  # Reset task section flag on new lesson
            if current_lesson:
                extracted_elements.append(current_lesson)
            lesson_title_match = re.match(r"(урок|lesson)\s*(\d+)\s*[:.-]?(.*)", text, re.IGNORECASE)
            if lesson_title_match:
                lesson_number = int(lesson_title_match.group(2))
                lesson_title = lesson_title_match.group(3).strip() or f"Lesson {lesson_number}"
            else:
                lesson_number = len(extracted_elements) + 1  # Fallback lesson number
                lesson_title = text
            current_lesson = {"lesson_number": lesson_number, "title": lesson_title, "content": []}
            print(f"DEBUG: New lesson at paragraph {i}: {lesson_number} - '{lesson_title}'")
        elif text:
            content_item = {"type": "explanation", "text": text}
            
            # Pattern 1: Check for diagram references using traditional pattern
            if DIAGRAM_PATTERN.search(text):
                content_item["type"] = "task"
                content_item["diagram_number_reference"] = int(DIAGRAM_PATTERN.search(text).group(2))
                task_count += 1
                print(f"DEBUG: Identified task type 1 at paragraph {i}: '{text}'")
            
            # Pattern 2: Check for dual task pattern (e.g., "25 _______  26 _______")
            elif DUAL_TASK_PATTERN.match(text):
                match = DUAL_TASK_PATTERN.match(text)
                
                # Create first task
                first_task = {"type": "task", "text": f"{match.group(1)} _______________"}
                first_task["diagram_number_reference"] = int(match.group(1))
                task_count += 1
                print(f"DEBUG: Identified dual task (first) at paragraph {i}: '{match.group(1)} _______________' - Diagram #{match.group(1)}")
                
                # Create second task
                second_task = {"type": "task", "text": f"{match.group(2)} _______________"}
                second_task["diagram_number_reference"] = int(match.group(2))
                task_count += 1
                print(f"DEBUG: Identified dual task (second) at paragraph {i}: '{match.group(2)} _______________' - Diagram #{match.group(2)}")
                
                # Add both tasks
                if current_lesson:
                    current_lesson["content"].append(first_task)
                    current_lesson["content"].append(second_task)
                else:
                    if not data["lessons"]:
                        default_lesson_title = "Introduction"
                        current_lesson = {"lesson_number": 0, "title": default_lesson_title, "content": []}
                        data["lessons"].append(current_lesson)
                    current_lesson["content"].append(first_task)
                    current_lesson["content"].append(second_task)
                
                # Skip further processing of this item
                continue
            
            # Pattern 3: Check for the single task pattern (e.g., "25 _______________")
            elif TASK_PATTERN.match(text) or TASK_PATTERN_SPACED.match(text):
                match = TASK_PATTERN.match(text) or TASK_PATTERN_SPACED.match(text)
                content_item["type"] = "task"
                diagram_number = int(match.group(1))
                content_item["diagram_number_reference"] = diagram_number
                task_count += 1
                print(f"DEBUG: Identified task type 3 at paragraph {i}: '{text}' - Diagram #{diagram_number}")
            
            # Pattern 4: If we're in a task section, consider numbered items as tasks
            elif in_task_section and NUMBER_PATTERN.match(text):
                content_item["type"] = "task"
                diagram_number = int(NUMBER_PATTERN.match(text).group(1))
                content_item["diagram_number_reference"] = diagram_number
                task_count += 1
                print(f"DEBUG: Identified task type 4 at paragraph {i}: '{text}' - Diagram #{diagram_number}")
            
            # NEW Pattern 5: Multi-number detection for tasks grouped on the same line
            elif in_task_section:
                # Look for multiple numbers in the text
                matches = MULTI_NUMBER_PATTERN.findall(text)
                if matches:
                    # Skip this paragraph but add individual task items for each number
                    print(f"DEBUG: Found multi-number task group at paragraph {i}: '{text}' with {len(matches)} numbers: {matches}")
                    for match in matches:
                        diagram_number = int(match)
                        new_task = {
                            "type": "general_task", 
                            "text": str(diagram_number),
                            "diagram_number_reference": diagram_number
                        }
                        if current_lesson:
                            current_lesson["content"].append(new_task)
                            task_count += 1
                            print(f"DEBUG: Created task from multi-number line for Diagram #{diagram_number}")
                    # Skip further processing of this paragraph
                    continue
            
            # Check for tasks with image in run
            has_image_in_run = False
            # Check if the paragraph itself contains an image
            for run in para.runs:
                if run._element.findall('.//w:drawing', docx.oxml.ns.nsmap) or \
                   run._element.findall('.//w:pict', docx.oxml.ns.nsmap):
                    has_image_in_run = True
                    break
            
            current_para_index = i # Save current paragraph index

            if has_image_in_run and content_item["type"] == "explanation":
                content_item["type"] = "task" # Assume it's a task if there's an image
                
                # Attempt to find diagram number in the next non-empty paragraph
                found_diagram_number_for_image = None
                # Look ahead up to 3 non-empty paragraphs or 5 total paragraphs
                non_empty_paras_checked = 0
                for next_para_offset in range(1, 6): # Check next 5 paragraphs
                    if current_para_index + next_para_offset < len(doc.paragraphs):
                        next_para_text = doc.paragraphs[current_para_index + next_para_offset].text.strip()
                        if next_para_text: # Found a non-empty paragraph
                            non_empty_paras_checked +=1
                            num_match = NUMBER_PATTERN.match(next_para_text)
                            if num_match:
                                found_diagram_number_for_image = int(num_match.group(1))
                                print(f"DEBUG: Found diagram number {found_diagram_number_for_image} in paragraph below image (para {current_para_index} -> {current_para_index + next_para_offset}) text: '{next_para_text}'")
                                break # Found number, stop
                            if non_empty_paras_checked >= 2: # Only check first 1-2 non-empty for number
                                break
                    else: # No more paragraphs
                        break
                
                if found_diagram_number_for_image is not None:
                    content_item["diagram_number_reference"] = found_diagram_number_for_image
                else:
                    # Fallback if no number found below
                    content_item["diagram_number_reference"] = task_count + 1 # Existing fallback, less ideal
                    print(f"DEBUG: No specific number found below image at para {current_para_index} ('{text}'), using fallback ref #{content_item['diagram_number_reference']}")
                
                task_count += 1
                print(f"DEBUG: Identified task with inline image at paragraph {current_para_index}: '{text}' - Assigned Diagram Ref #{content_item['diagram_number_reference']}")
            
            if current_lesson:
                current_lesson["content"].append(content_item)
            else:  # Content before any lesson detected
                if not data["lessons"]:  # Create a default lesson if none exists
                    default_lesson_title = "Introduction"
                    current_lesson = {"lesson_number": 0, "title": default_lesson_title, "content": []}
                    data["lessons"].append(current_lesson)
                current_lesson["content"].append(content_item)

    # Extract images (runs and inline_shapes)
    # This part is complex as images can be in different places in DOCX
    # and direct mapping to paragraphs is not straightforward via python-docx.
    # We'll extract all images and try to map them based on diagram numbers later.
    
    # Attempt to extract images associated with runs (more common for inline)
    img_idx = 0
    img_count = 0
    
    # Collect diagram references to align image numbers with diagram numbers
    diagram_refs = []
    for lesson in data["lessons"] + ([current_lesson] if current_lesson else []):
        for item in lesson.get("content", []):
            if (item.get("type") == "task" or item.get("type") == "general_task") and "diagram_number_reference" in item:
                diagram_refs.append(item["diagram_number_reference"])
    
    # Sort diagram numbers to establish sequential order and remove duplicates
    diagram_refs = sorted(set(diagram_refs))
    print(f"DEBUG: Found {len(diagram_refs)} unique diagram references: {diagram_refs}...")
    
    # Count images in the document
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            img_count += 1
            
    print(f"DEBUG: Document has {img_count} image relations")
    
    # Start image counter at 1 if no diagram references found
    next_image_number = 1
    
    # Extract and process images
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_bytes = rel.target_part.blob
            try:
                img = Image.open(io.BytesIO(image_bytes))
                
                # If we can align with known diagram numbers, use those
                # Otherwise use sequential numbering
                if img_idx < len(diagram_refs):
                    # Use diagrams refs for image numbers to align with tasks
                    image_number = diagram_refs[img_idx]
                else:
                    # If we've gone through all known diagram refs, use sequential numbering 
                    # starting after the highest known diagram number
                    image_number = next_image_number
                    next_image_number += 1
                
                # Create filename with the aligned image number
                img_filename = f"diagram_doc_img{image_number}.{img.format.lower() or 'png'}"
                img_path = os.path.join(output_images_dir, img_filename)
                img.save(img_path)
                img_idx += 1
                print(f"DEBUG: Saved image {img_idx}/{img_count}: {img_filename} with number {image_number}")
                
                # Use improved image-task association logic
                if current_lesson:
                    # Try to associate image with task based on image number
                    if not associate_image_with_task(data["lessons"] + [current_lesson], img_filename, image_number):
                        print(f"DEBUG: Could not find a task to associate with image {img_filename}")
                else:
                    print(f"DEBUG: Could not find a task to associate with image {img_filename}")
            except Exception as e:
                print(f"Warning: Could not process image {rel.target_ref}: {e}")

    if current_lesson:
        extracted_elements.append(current_lesson)
    
    data["lessons"] = extracted_elements
    print(f"DEBUG: Extraction complete. Found {task_count} tasks and {img_count} images.")
    return data, []  # Returning empty list for images for now, as they are linked in content

def extract_text_and_images_from_pdf(file_path: str, output_images_dir: str) -> Tuple[Dict[str, Any], List[Dict[str, Any]]]:
    """Extracts text and images from a PDF file."""
    try:
        doc = fitz.open(file_path)
        print(f"Successfully opened PDF file: {file_path}")
        print(f"Document has {len(doc)} pages")
    except Exception as e:
        print(f"Error opening PDF file: {e}")
        return {"book_title": "Error", "lessons": []}, []
    
    data = {"book_title": "Unknown PDF Document", "lessons": []}
    current_lesson = None
    image_counter = 1  # Unique counter for PDF images
    task_count = 0
    
    # Track if we're in a task section to improve task identification
    in_task_section = False
    
    # Extract metadata for title if available
    if doc.metadata and doc.metadata.get("title"):
        data["book_title"] = doc.metadata["title"]
    else:
        data["book_title"] = os.path.splitext(os.path.basename(file_path))[0]

    # Consolidate extracted elements (lessons)
    extracted_lessons = []

    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        text_blocks = page.get_text("dict")["blocks"]
        
        for block in text_blocks:
            if block["type"] == 0:  # Text block
                for line in block["lines"]:
                    line_text = "".join([span["text"] for span in line["spans"]]).strip()
                    if not line_text:
                        continue
                        
                    # Detect task sections
                    if TASK_SECTION_PATTERN.search(line_text):
                        in_task_section = True
                        print(f"DEBUG: Entering task section: '{line_text}'")
                    
                    # Crude heading/lesson detection based on font size or keywords
                    # This is highly dependent on PDF structure and needs refinement
                    is_heading = False
                    if line["spans"]:
                        span_font_size = line["spans"][0]["size"]
                        # Example: consider text as heading if font size > 15
                        if span_font_size > 15 or line_text.lower().startswith("урок") or line_text.lower().startswith("lesson"):
                            is_heading = True
                            in_task_section = False  # Reset task section flag on new lesson

                    if is_heading:
                        if current_lesson:
                            extracted_lessons.append(current_lesson)
                        
                        lesson_title_match = re.match(r"(урок|lesson)\s*(\d+)\s*[:.-]?(.*)", line_text, re.IGNORECASE)
                        if lesson_title_match:
                            lesson_number = int(lesson_title_match.group(2))
                            lesson_title = lesson_title_match.group(3).strip() or f"Lesson {lesson_number}"
                        else:
                            lesson_number = len(extracted_lessons) + 1
                            lesson_title = line_text
                        
                        current_lesson = {"lesson_number": lesson_number, "title": lesson_title, "content": []}
                        print(f"DEBUG: New lesson: {lesson_number} - '{lesson_title}'")
                        continue  # Move to next line after processing heading

                    if line_text:
                        content_item = {"type": "explanation", "text": line_text}
                        
                        # Pattern 1: Check for diagram references using traditional pattern
                        if DIAGRAM_PATTERN.search(line_text):
                            content_item["type"] = "task"
                            content_item["diagram_number_reference"] = int(DIAGRAM_PATTERN.search(line_text).group(2))
                            task_count += 1
                            print(f"DEBUG: Identified task type 1: '{line_text}'")
                        
                        # Pattern 2: Check for dual task pattern (e.g., "25 _______  26 _______")
                        elif DUAL_TASK_PATTERN.match(line_text):
                            match = DUAL_TASK_PATTERN.match(line_text)
                            
                            # Create first task
                            first_task = {"type": "task", "text": f"{match.group(1)} _______________"}
                            first_task["diagram_number_reference"] = int(match.group(1))
                            task_count += 1
                            print(f"DEBUG: Identified dual task (first): '{match.group(1)} _______________' - Diagram #{match.group(1)}")
                            
                            # Create second task
                            second_task = {"type": "task", "text": f"{match.group(2)} _______________"}
                            second_task["diagram_number_reference"] = int(match.group(2))
                            task_count += 1
                            print(f"DEBUG: Identified dual task (second): '{match.group(2)} _______________' - Diagram #{match.group(2)}")
                            
                            # Add both tasks
                            if current_lesson:
                                current_lesson["content"].append(first_task)
                                current_lesson["content"].append(second_task)
                            else:
                                if not extracted_lessons:
                                    default_lesson_title = data["book_title"] + " - Introduction"
                                    current_lesson = {"lesson_number": 0, "title": default_lesson_title, "content": []}
                                    extracted_lessons.append(current_lesson)
                                current_lesson["content"].append(first_task)
                                current_lesson["content"].append(second_task)
                            
                            # Skip further processing of this item
                            continue
                        
                        # Pattern 3: Check for the single task pattern (e.g., "25 _______________")
                        elif TASK_PATTERN.match(line_text) or TASK_PATTERN_SPACED.match(line_text):
                            match = TASK_PATTERN.match(line_text) or TASK_PATTERN_SPACED.match(line_text)
                            content_item["type"] = "task"
                            diagram_number = int(match.group(1))
                            content_item["diagram_number_reference"] = diagram_number
                            task_count += 1
                            print(f"DEBUG: Identified task type 3: '{line_text}' - Diagram #{diagram_number}")
                        
                        # Pattern 4: If we're in a task section, consider numbered items as tasks
                        elif in_task_section and NUMBER_PATTERN.match(line_text):
                            content_item["type"] = "task"
                            diagram_number = int(NUMBER_PATTERN.match(line_text).group(1))
                            content_item["diagram_number_reference"] = diagram_number
                            task_count += 1
                            print(f"DEBUG: Identified task type 4: '{line_text}' - Diagram #{diagram_number}")
                        
                        if current_lesson:
                            current_lesson["content"].append(content_item)
                        else:  # Content before any lesson detected
                            if not extracted_lessons:
                                default_lesson_title = data["book_title"] + " - Introduction"
                                current_lesson = {"lesson_number": 0, "title": default_lesson_title, "content": []}
                                extracted_lessons.append(current_lesson)
                            current_lesson["content"].append(content_item)
        
        # Extract images from the page
        img_list = page.get_images(full=True)
        
        # Collect diagram references to align image numbers with diagram numbers
        if page_num == 0:  # Only do this once at the start
            diagram_refs = []
            for lesson in extracted_lessons + ([current_lesson] if current_lesson else []):
                for item in lesson.get("content", []):
                    if item.get("type") == "task" and "diagram_number_reference" in item:
                        diagram_refs.append(item["diagram_number_reference"])
            
            # Sort diagram numbers to establish sequential order
            diagram_refs = sorted(set(diagram_refs))
            print(f"DEBUG: Found {len(diagram_refs)} unique diagram references: {diagram_refs[:10]}...")
            
            # Start image counter after highest known diagram reference
            next_image_number = 1
            
        for img_index, img_info in enumerate(img_list):
            xref = img_info[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            image_ext = base_image["ext"]
            
            # If we can align with known diagram numbers, use those
            # Otherwise use sequential numbering
            if img_index < len(diagram_refs):
                # Use diagram refs for image numbers to align with tasks
                image_number = diagram_refs[img_index]
            else:
                # If we've gone through all known diagram refs, use sequential numbering
                image_number = next_image_number
                next_image_number += 1
                
            img_filename = f"diagram_pdf_page{page_num+1}_img{image_number}.{image_ext}"
            img_path = os.path.join(output_images_dir, img_filename)
            
            with open(img_path, "wb") as img_file:
                img_file.write(image_bytes)
                
            print(f"DEBUG: Saved PDF image {img_filename} with number {image_number}")
            
            # Use improved image-task association logic
            if not associate_image_with_task(extracted_lessons + ([current_lesson] if current_lesson else []), 
                                        img_filename, image_number):
                print(f"DEBUG: Could not find a task to associate with image {img_filename}")
            
        # Increment overall image counter for the next page
        image_counter += len(img_list)

    if current_lesson:  # Add the last processed lesson
        extracted_lessons.append(current_lesson)
    
    data["lessons"] = extracted_lessons
    print(f"DEBUG: Extraction complete. Found {task_count} tasks across {len(extracted_lessons)} lessons.")
    return data, []

def extract_content(file_path: str, output_images_dir: str = None) -> Dict[str, Any]:
    """
    Extracts content from a document file (DOCX or PDF) and saves any images.
    
    Args:
        file_path: Path to the document file
        output_images_dir: Directory to save extracted images (optional)
    
    Returns:
        Dictionary containing extracted text and image information
    """
    if output_images_dir is None:
        output_images_dir = config.OUTPUT_IMAGE_DIR
        
    file_extension = os.path.splitext(file_path)[1].lower()
    
    if file_extension == '.docx':
        extracted_data, images_info = extract_text_and_images_from_docx(file_path, output_images_dir)
    elif file_extension == '.pdf':
        extracted_data, images_info = extract_text_and_images_from_pdf(file_path, output_images_dir)
    else:
        raise ValueError(f"Unsupported file format: {file_extension}")
    
    print(f"DEBUG: Extracted {len(images_info)} images from document")
    
    # Apply enhanced diagram processing if available
    if ENHANCED_DIAGRAM_PROCESSOR_AVAILABLE and extracted_data and images_info:
        try:
            print("Applying enhanced diagram association processing...")
            extracted_data = apply_enhanced_diagram_associations(extracted_data, images_info)
            print("Enhanced diagram processing completed successfully")
        except Exception as e:
            print(f"Enhanced diagram processing failed, falling back to legacy: {e}")
            # Fall back to legacy association logic
            _legacy_associate_images_with_tasks(extracted_data, images_info)
    else:
        # Use legacy association logic
        print("Using legacy diagram association logic")
        _legacy_associate_images_with_tasks(extracted_data, images_info)
    
    return extracted_data

def _legacy_associate_images_with_tasks(extracted_data: Dict[str, Any], images_info: List[Dict[str, Any]]) -> None:
    """Legacy method for associating images with tasks"""
    # This contains the original association logic from the current extract.py
    print("Attempting to associate images with tasks using legacy method...")
    
    lessons = extracted_data.get("lessons", [])
    
    for img_info in images_info:
        img_filename = img_info.get('filename', '')
        
        # Try to associate each image with a task using the existing logic
        success = associate_image_with_task(lessons, img_filename)
        
        if success:
            print(f"Successfully associated image {img_filename} with a task (legacy method)")
        else:
            print(f"Could not associate image {img_filename} with any task (legacy method)")

if __name__ == "__main__":
    # Example usage
    test_file = "test_document.docx"  # Replace with your test document
    output_dir = "output_images"
    os.makedirs(output_dir, exist_ok=True)
    
    result = extract_content(test_file, output_dir)
    if result:
        print(json.dumps(result, indent=2))
    else:
        print("Extraction failed.") 